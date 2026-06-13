#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera o planejamento aula-por-aula (Model UN) em Word (.docx) com layout
caprichado e identidade visual da Aluizio Educação.

Uso:
    python3 scripts/gerar_planejamento_mun.py [caminho_logo.png]

Se um caminho de logo for passado e o arquivo existir, ele é embutido no
cabeçalho; caso contrário, um wordmark textual "Aluizio Educação" é usado
como espaço reservado (o .docx é editável e a logo pode ser trocada).
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Paleta (diplomática / acadêmica): navy + azul ONU + dourado
NAVY   = "1B2A4A"
BLUE   = "2E5AAC"
GOLD   = "C9A227"
LIGHT  = "EEF2FA"
LIGHT2 = "F7F9FC"
GRAY   = "3C4250"
WHITE  = "FFFFFF"

FONT = "Calibri"
DISPLAY = "Georgia"

# ----------------------------------------------------------------------------
# Helpers de baixo nível (XML)

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left),
                     ("end", right), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)

def _border_edge(elm, edge, sz, color, val="single"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    elm.append(e)

def table_borders(table, sz=4, color="D7DEEA", inside=True):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        _border_edge(borders, edge, sz, color)
    tblPr.append(borders)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)

def set_width(table, total_inches, col_fractions):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(total_inches * col_fractions[i])

def runfmt(run, size=10.5, color=GRAY, bold=False, italic=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    # garante a fonte também para east-asian/cs
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)

def para_spacing(p, before=4, after=4, line=1.12):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def cell_text(cell, text, size=10.5, color=GRAY, bold=False, align=None,
              italic=False, font=FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    para_spacing(p, 2, 2, 1.05)
    r = p.add_run(text)
    runfmt(r, size, color, bold, italic, font)
    return p

# ----------------------------------------------------------------------------
# Blocos de alto nível

def add_logo_header(doc, logo_path=None):
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)
    set_width(t, 6.5, [0.62, 0.38])
    left, right = t.rows[0].cells
    shade(left, NAVY); shade(right, NAVY)
    set_cell_margins(left, 140, 140, 200, 120)
    set_cell_margins(right, 140, 140, 120, 200)
    vcenter(left); vcenter(right)

    # Esquerda: logo (imagem) ou wordmark
    left.text = ""
    p = left.paragraphs[0]
    para_spacing(p, 0, 0, 1.0)
    if logo_path:
        try:
            r = p.add_run()
            r.add_picture(logo_path, height=Inches(0.55))
        except Exception:
            logo_path = None
    if not logo_path:
        r = p.add_run("Aluizio")
        runfmt(r, 22, WHITE, bold=True, font=DISPLAY)
        r2 = p.add_run("  Educação")
        runfmt(r2, 22, GOLD, bold=True, font=DISPLAY)
        p2 = left.add_paragraph()
        para_spacing(p2, 2, 0, 1.0)
        r3 = p2.add_run("Mentoria & preparação acadêmica")
        runfmt(r3, 9, "C7D0E0", italic=True)

    # Direita: etiqueta do documento
    right.text = ""
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_spacing(pr, 0, 0, 1.0)
    r = pr.add_run("PLANO DE CURSO")
    runfmt(r, 10, GOLD, bold=True)
    pr2 = right.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_spacing(pr2, 2, 0, 1.0)
    r = pr2.add_run("10 horas · 5 encontros")
    runfmt(r, 9, "C7D0E0")

def add_title_block(doc):
    p = doc.add_paragraph()
    para_spacing(p, 14, 0, 1.05)
    r = p.add_run("Geopolítica Aplicada às Simulações Model UN")
    runfmt(r, 22, NAVY, bold=True, font=DISPLAY)
    p2 = doc.add_paragraph()
    para_spacing(p2, 2, 6, 1.15)
    r = p2.add_run("Planejamento aula por aula — dominar as estruturas da ONU "
                   "e aplicá-las para resolver problemas das simulações "
                   "(negociação, redação de resoluções e obtenção de fundos).")
    runfmt(r, 11.5, GRAY, italic=True)
    # linha dourada
    line = doc.add_table(rows=1, cols=1)
    no_borders(line); set_width(line, 6.5, [1.0])
    shade(line.rows[0].cells[0], GOLD)
    line.rows[0].cells[0].height = Pt(3)
    cell_text(line.rows[0].cells[0], "", size=2)

def add_meta_strip(doc):
    items = [("Carga horária", "10 horas"),
             ("Formato", "5 encontros de 2h"),
             ("Modalidade", "Individual"),
             ("Público", "Estudante MUN")]
    t = doc.add_table(rows=1, cols=4)
    no_borders(t); set_width(t, 6.5, [0.25]*4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(items):
        c = t.rows[0].cells[i]
        shade(c, LIGHT); set_cell_margins(c, 90, 90, 120, 120)
        c.text = ""
        p = c.paragraphs[0]; para_spacing(p, 0, 0, 1.0)
        r = p.add_run(k.upper()); runfmt(r, 7.5, BLUE, bold=True)
        p2 = c.add_paragraph(); para_spacing(p2, 1, 0, 1.0)
        r = p2.add_run(v); runfmt(r, 11, NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def section_bar(doc, text, sub=None):
    t = doc.add_table(rows=1, cols=1)
    no_borders(t); set_width(t, 6.5, [1.0])
    c = t.rows[0].cells[0]
    shade(c, BLUE); set_cell_margins(c, 70, 70, 140, 140)
    c.text = ""
    p = c.paragraphs[0]; para_spacing(p, 0, 0, 1.0)
    r = p.add_run(text); runfmt(r, 12, WHITE, bold=True)
    if sub:
        r2 = p.add_run("    " + sub); runfmt(r2, 9.5, "DCE6FB")
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

def body(doc, text, size=10.5, before=2, after=4, bold=False, color=GRAY):
    p = doc.add_paragraph()
    para_spacing(p, before, after, 1.18)
    r = p.add_run(text); runfmt(r, size, color, bold=bold)
    return p

def bullet(doc, text, label=None):
    p = doc.add_paragraph()
    para_spacing(p, 1, 1, 1.16)
    p.paragraph_format.left_indent = Inches(0.18)
    rb = p.add_run("•  "); runfmt(rb, 10.5, GOLD, bold=True)
    if label:
        rl = p.add_run(label + " "); runfmt(rl, 10.5, NAVY, bold=True)
    r = p.add_run(text); runfmt(r, 10.5, GRAY)
    return p

def schedule_table(doc, rows):
    t = doc.add_table(rows=1, cols=3)
    table_borders(t, 4, "D7DEEA")
    set_width(t, 6.5, [0.13, 0.62, 0.25])
    hdr = t.rows[0].cells
    for i, h in enumerate(["Aula", "Tema central", "Carga"]):
        shade(hdr[i], NAVY); set_cell_margins(hdr[i], 60, 60, 110, 110)
        cell_text(hdr[i], h, 10, WHITE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    for idx, (num, tema, carga) in enumerate(rows):
        row = t.add_row().cells
        bg = WHITE if idx % 2 == 0 else LIGHT2
        for c in row:
            shade(c, bg); set_cell_margins(c, 60, 60, 110, 110); vcenter(c)
        cell_text(row[0], num, 11, BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[1], tema, 10.5, GRAY)
        cell_text(row[2], carga, 10, GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

def lesson_header(doc, num, titulo, dur="2 horas (120 min)"):
    t = doc.add_table(rows=1, cols=2)
    no_borders(t); set_width(t, 6.5, [0.78, 0.22])
    l, r = t.rows[0].cells
    shade(l, NAVY); shade(r, GOLD)
    set_cell_margins(l, 80, 80, 140, 100); set_cell_margins(r, 80, 80, 100, 140)
    vcenter(l); vcenter(r)
    l.text = ""
    p = l.paragraphs[0]; para_spacing(p, 0, 0, 1.0)
    ra = p.add_run(f"AULA {num}  ·  "); runfmt(ra, 11, GOLD, bold=True)
    rt = p.add_run(titulo); runfmt(rt, 12, WHITE, bold=True)
    cell_text(r, dur, 9.5, NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(1)

def objetivo(doc, texto):
    t = doc.add_table(rows=1, cols=1)
    no_borders(t); set_width(t, 6.5, [1.0])
    c = t.rows[0].cells[0]
    shade(c, LIGHT); set_cell_margins(c, 70, 70, 140, 140)
    c.text = ""
    p = c.paragraphs[0]; para_spacing(p, 0, 0, 1.12)
    r = p.add_run("Objetivo da aula  "); runfmt(r, 9.5, BLUE, bold=True)
    r2 = p.add_run(texto); runfmt(r2, 10, GRAY)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(1)

def timeblocks(doc, blocks):
    label = doc.add_paragraph(); para_spacing(label, 4, 2, 1.0)
    r = label.add_run("Roteiro de tempo"); runfmt(r, 9.5, NAVY, bold=True)
    t = doc.add_table(rows=1, cols=2)
    table_borders(t, 4, "D7DEEA")
    set_width(t, 6.5, [0.2, 0.8])
    hdr = t.rows[0].cells
    for i, h in enumerate(["Tempo", "Atividade"]):
        shade(hdr[i], BLUE); set_cell_margins(hdr[i], 50, 50, 110, 110)
        cell_text(hdr[i], h, 9.5, WHITE, bold=True)
    for idx, (tempo, atv) in enumerate(blocks):
        row = t.add_row().cells
        bg = WHITE if idx % 2 == 0 else LIGHT2
        for c in row:
            shade(c, bg); set_cell_margins(c, 50, 50, 110, 110); vcenter(c)
        cell_text(row[0], tempo, 9.5, BLUE, bold=True)
        cell_text(row[1], atv, 10, GRAY)

def callout(doc, label, texto, color=GOLD):
    t = doc.add_table(rows=1, cols=1)
    table_borders(t, 4, color)
    set_width(t, 6.5, [1.0])
    c = t.rows[0].cells[0]
    shade(c, LIGHT2); set_cell_margins(c, 70, 70, 140, 140)
    c.text = ""
    p = c.paragraphs[0]; para_spacing(p, 0, 0, 1.12)
    r = p.add_run(f"{label}  "); runfmt(r, 9.5, color if color != GOLD else "9A7B16", bold=True)
    r2 = p.add_run(texto); runfmt(r2, 10, GRAY)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)

def spacer(doc, pt=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pt)

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aluizio Educação  ·  [e-mail]  ·  [telefone/WhatsApp]  ·  [site]   —   página ")
    runfmt(r, 8, "9AA3B2")
    # campo de número de página
    fldStart = OxmlElement('w:fldSimple'); fldStart.set(qn('w:instr'), 'PAGE')
    run_holder = p.add_run()
    run_holder._element.addprevious(fldStart)
    r2 = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '16'); rpr.append(sz)
    r2.append(rpr); t = OxmlElement('w:t'); t.text = '1'; r2.append(t)
    fldStart.append(r2)

# ----------------------------------------------------------------------------
# CONTEÚDO

CRONOGRAMA = [
    ("1", "O sistema ONU: arquitetura e lógica de funcionamento", "2h"),
    ("2", "Órgãos, agências e mecanismos de financiamento", "2h"),
    ("3", "Regras de procedimento e dinâmica de comitê", "2h"),
    ("4", "Documentos, redação de resoluções e negociação", "2h"),
    ("5", "Simulação integrada, oratória e feedback", "2h"),
]

AULAS = [
 {
  "num": "1", "titulo": "O sistema ONU: arquitetura e lógica de funcionamento",
  "objetivo": "Compreender a origem, os princípios e os órgãos da ONU e como cada "
              "um se traduz em um comitê de simulação.",
  "blocks": [
     ("0–10 min", "Abertura, diagnóstico do aluno e objetivos da aula."),
     ("10–40 min", "A Carta de São Francisco (1945): propósitos e princípios — soberania, "
                   "não-intervenção, segurança coletiva e solução pacífica de controvérsias."),
     ("40–80 min", "Os órgãos principais e suas competências: Assembleia Geral, Conselho de "
                   "Segurança (P5 e veto), ECOSOC, Secretariado, Corte Internacional de Justiça "
                   "e o Conselho de Tutela."),
     ("80–105 min", "Da estrutura ao comitê: como cada órgão vira um comitê de MUN e o que muda "
                    "em poderes e documentos."),
     ("105–120 min", "Atividade prática e fechamento."),
  ],
  "pratica": "Mapa do sistema ONU — dado um conjunto de problemas (conflito armado, epidemia, "
             "refugiados, crise alimentar), o aluno indica o órgão/comitê competente e justifica.",
  "entregavel": "Quadro-resumo “qual órgão para qual problema”.",
 },
 {
  "num": "2", "titulo": "Órgãos, agências e mecanismos de financiamento",
  "objetivo": "Distinguir órgãos, programas, fundos e agências, entender como a ONU se "
              "financia e como obter recursos para propostas em comitê.",
  "blocks": [
     ("0–10 min", "Revisão da aula anterior e conexão com o tema."),
     ("10–35 min", "O ecossistema multilateral: programas e fundos (PNUD, UNICEF, ACNUR, "
                   "PMA/WFP) × agências especializadas (OMS, FAO, UNESCO, OIT) e Bretton Woods."),
     ("35–75 min", "Como a ONU se financia: contribuições obrigatórias × voluntárias, escala "
                   "de cotas, fundos fiduciários (trust funds) e pooled funds (ex.: CERF)."),
     ("75–105 min", "Obtendo fundos na simulação: identificar a fonte certa, redigir cláusula "
                    "que solicita/cria financiamento, mobilizar doadores e prever prestação de contas."),
     ("105–120 min", "Atividade prática e fechamento."),
  ],
  "pratica": "Captação de recursos — diante de uma crise humanitária fictícia, o aluno desenha "
             "o plano de financiamento (fundos/agências a acionar, cláusula a redigir, parceiros).",
  "entregavel": "Rascunho de cláusula operativa de financiamento.",
 },
 {
  "num": "3", "titulo": "Regras de procedimento e dinâmica de comitê",
  "objetivo": "Dominar as Rules of Procedure e o fluxo de uma sessão para atuar com "
              "segurança na mesa.",
  "blocks": [
     ("0–10 min", "Revisão e aquecimento."),
     ("10–45 min", "Rules of Procedure: quórum, lista de oradores (GSL), moções (caucus "
                   "moderado/não-moderado, encerrar debate, votação) e pontos (ordem, informação)."),
     ("45–75 min", "Fluxo da sessão: roll call → setting the agenda → debate formal → caucus → "
                   "redação → votação."),
     ("75–100 min", "Papéis e relação com a mesa (dais/chair); etiqueta; tipos de comitê "
                    "(AGNU, CSNU, ECOSOC, históricos e de crise)."),
     ("100–120 min", "Mini-simulação de procedimento e fechamento."),
  ],
  "pratica": "Mini-simulação: o aluno conduz uma rodada com abertura de agenda, moções de "
             "caucus e votação.",
  "entregavel": "“Cola” pessoal de moções (quando usar cada uma).",
 },
 {
  "num": "4", "titulo": "Documentos, redação de resoluções e negociação",
  "objetivo": "Redigir documentos tecnicamente corretos e negociar para construir consenso "
              "em torno de uma resolução.",
  "blocks": [
     ("0–10 min", "Revisão e exemplos comentados."),
     ("10–35 min", "Documentos do delegado: position paper, working paper, draft resolution "
                   "e emendas."),
     ("35–70 min", "Anatomia da resolução: cláusulas preambulares × operativas e os verbos "
                   "típicos de cada uma."),
     ("70–95 min", "Cláusula operativa exequível: criar mecanismos, solicitar financiamento, "
                   "definir prazos, responsáveis e monitoramento."),
     ("95–120 min", "Redação assistida e negociação simulada."),
  ],
  "pratica": "O aluno redige uma draft resolution completa — incluindo um mecanismo de "
             "financiamento — e a defende em negociação simulada (blocos, sponsors, concessões).",
  "entregavel": "Draft resolution revisada com feedback.",
 },
 {
  "num": "5", "titulo": "Simulação integrada, oratória e feedback",
  "objetivo": "Integrar tudo numa simulação completa, com oratória diplomática e resposta a "
              "problemas-surpresa, incluindo crises de financiamento.",
  "blocks": [
     ("0–15 min", "Oratória diplomática: discurso de abertura, intervenções na GSL, direito "
                  "de resposta, linguagem formal e controle do tempo."),
     ("15–30 min", "Postura, protocolo e funcionamento dos comitês de crise."),
     ("30–100 min", "Simulação completa cronometrada: cenário com um problema a resolver de "
                    "ponta a ponta — da abertura à votação de uma resolução que mobiliza recursos."),
     ("100–115 min", "Feedback individualizado por rubrica."),
     ("115–120 min", "Plano de evolução pós-curso e encerramento."),
  ],
  "pratica": "Simulação integrada cronometrada, avaliada por rubrica (procedimento, conteúdo, "
             "negociação, oratória e documentos).",
  "entregavel": "Ficha de feedback + plano de estudos personalizado.",
 },
]


def build(logo_path=None, out="docs/cursos/Planejamento_MUN_AluizioEducacao.docx"):
    doc = Document()
    # margens
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    # estilo base
    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(GRAY)

    add_logo_header(doc, logo_path)
    add_title_block(doc)
    add_meta_strip(doc)

    section_bar(doc, "Apresentação")
    body(doc, "Este programa prepara o estudante para atuar com estratégia em simulações "
              "Model UN. Mais do que conhecer o organograma da ONU, o objetivo é que o aluno "
              "entenda a lógica das estruturas das Nações Unidas e saiba acioná-las para "
              "resolver problemas concretos da simulação — escolher o órgão competente, "
              "negociar, redigir resoluções e, sobretudo, mobilizar e obter fundos para "
              "viabilizar suas propostas. O fio condutor é “da estrutura à ação”.")

    section_bar(doc, "Objetivos de aprendizagem")
    for b in [
        "Identificar os órgãos da ONU, suas competências e seus limites.",
        "Distinguir órgãos, programas, fundos e agências — e saber a qual recorrer.",
        "Compreender como a ONU se financia e como obter recursos em comitê.",
        "Dominar as regras de procedimento (moções, caucus, votação).",
        "Redigir position paper, working paper e draft resolution.",
        "Negociar, formar blocos e construir consenso, inclusive sob crise.",
        "Apresentar-se com postura e oratória diplomáticas.",
    ]:
        bullet(doc, b)

    spacer(doc, 4)
    section_bar(doc, "Cronograma das 10 horas")
    schedule_table(doc, CRONOGRAMA)

    # quebra de página antes do detalhamento
    doc.add_page_break()
    section_bar(doc, "Planejamento aula por aula", "10 horas · 5 encontros de 2h")
    spacer(doc, 2)

    for i, a in enumerate(AULAS):
        lesson_header(doc, a["num"], a["titulo"])
        objetivo(doc, a["objetivo"])
        timeblocks(doc, a["blocks"])
        spacer(doc, 2)
        callout(doc, "Atividade prática", a["pratica"], GOLD)
        callout(doc, "Entregável do aluno", a["entregavel"], BLUE)
        if i < len(AULAS) - 1:
            doc.add_page_break()

    # Metodologia / avaliação
    doc.add_page_break()
    section_bar(doc, "Metodologia, materiais e avaliação")
    body(doc, "Metodologia.", bold=True, after=1)
    for b in [
        "Aprendizagem ativa: cada conceito é seguido de aplicação prática.",
        "Estudos de caso reais (operações de paz, apelos humanitários, resoluções históricas).",
        "Simulações progressivas, de micro-exercícios à simulação integrada final.",
        "Feedback contínuo por rubrica (procedimento, conteúdo, negociação, oratória, documentos).",
    ]:
        bullet(doc, b)
    spacer(doc, 3)
    body(doc, "Materiais.", bold=True, after=1)
    for b in [
        "Carta das Nações Unidas e regimento-modelo de MUN (Rules of Procedure).",
        "Modelos comentados de position paper, working paper e draft resolution.",
        "Glossário ONU/MUN (PT-BR / EN) e lista de fontes oficiais de pesquisa.",
    ]:
        bullet(doc, b)
    spacer(doc, 3)
    body(doc, "Avaliação.", bold=True, after=1)
    for b in [
        "Formativa: desempenho nos exercícios e entregáveis de cada aula.",
        "Somativa: rubrica aplicada na simulação integrada final (Aula 5).",
    ]:
        bullet(doc, b)

    spacer(doc, 6)
    callout(doc, "Resultado esperado",
            "Ao concluir, o estudante chega à simulação sabendo não apenas o que é a ONU, "
            "mas como usá-la: escolhe o órgão certo, domina o procedimento, redige documentos "
            "sólidos, negocia com método e sabe mobilizar e obter recursos para suas propostas.",
            NAVY)

    add_footer(doc)
    doc.save(out)
    print("Documento salvo em:", out)


if __name__ == "__main__":
    logo = sys.argv[1] if len(sys.argv) > 1 else None
    build(logo)
