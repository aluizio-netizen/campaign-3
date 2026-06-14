#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera a PROPOSTA COMERCIAL (Model UN) em Word (.docx), no mesmo padrão visual
do planejamento, com a identidade da Aluizio Educação.

Uso:
    python3 scripts/gerar_proposta_mun.py [caminho_logo.png]
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import gerar_planejamento_mun as base
from gerar_planejamento_mun import (
    NAVY, BLUE, GOLD, LIGHT, LIGHT2, GRAY, WHITE, FONT, DISPLAY,
    shade, set_cell_margins, vcenter, table_borders, no_borders, set_width,
    runfmt, para_spacing, cell_text, body, bullet, callout, spacer, section_bar,
    brand_header,
)


def title_block(doc):
    p = doc.add_paragraph(); para_spacing(p, 14, 0, 1.05)
    r = p.add_run("Mentoria de Preparação para Simulações Model UN")
    runfmt(r, 21, NAVY, bold=True, font=DISPLAY)
    p2 = doc.add_paragraph(); para_spacing(p2, 2, 6, 1.15)
    r = p2.add_run("Programa individual de 10 horas para dominar as estruturas da ONU "
                   "e aplicá-las na resolução de problemas das simulações.")
    runfmt(r, 11.5, GRAY, italic=True)
    line = doc.add_table(rows=1, cols=1); no_borders(line); set_width(line, 6.5, [1.0])
    shade(line.rows[0].cells[0], GOLD); cell_text(line.rows[0].cells[0], "", size=2)


def info_block(doc):
    pairs = [("De", "[Nome do educador / marca] — [e-mail] — [telefone/WhatsApp]"),
             ("Para", "[Nome do aluno / responsável]"),
             ("Data", "[data]"),
             ("Validade", "15 dias a partir da emissão")]
    t = doc.add_table(rows=len(pairs), cols=2)
    no_borders(t); set_width(t, 6.5, [0.16, 0.84])
    for i, (k, v) in enumerate(pairs):
        kc, vc = t.rows[i].cells
        shade(kc, LIGHT if i % 2 == 0 else LIGHT2)
        shade(vc, LIGHT if i % 2 == 0 else LIGHT2)
        set_cell_margins(kc, 50, 50, 110, 80); set_cell_margins(vc, 50, 50, 80, 110)
        vcenter(kc); vcenter(vc)
        cell_text(kc, k.upper(), 8.5, BLUE, bold=True)
        cell_text(vc, v, 10, GRAY)
    spacer(doc, 4)


def invest_table(doc):
    rows = [("Hora-aula", "R$ [200]"),
            ("Pacote 10 horas", "R$ [2.000]"),
            ("Pacote 10 horas à vista (desconto [10%])", "R$ [1.800]")]
    t = doc.add_table(rows=1, cols=2)
    table_borders(t, 4, "D7DEEA"); set_width(t, 6.5, [0.7, 0.3])
    h = t.rows[0].cells
    for i, txt in enumerate(["Mentoria individual (1 aluno)", "Valor"]):
        shade(h[i], NAVY); set_cell_margins(h[i], 60, 60, 110, 110)
        cell_text(h[i], txt, 10, WHITE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.RIGHT if i == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    for idx, (k, v) in enumerate(rows):
        row = t.add_row().cells
        bg = WHITE if idx % 2 == 0 else LIGHT2
        for c in row:
            shade(c, bg); set_cell_margins(c, 60, 60, 110, 110); vcenter(c)
        destaque = idx >= 1
        cell_text(row[0], k, 10.5, GRAY, bold=False)
        cell_text(row[1], v, 11 if destaque else 10.5, NAVY if destaque else GRAY,
                  bold=destaque, align=WD_ALIGN_PARAGRAPH.RIGHT)


def cenarios_table(doc):
    rows = [("Individual", "R$ [2.000]", "Atenção 100% personalizada"),
            ("Dupla (2 alunos)", "R$ [1.300]/aluno", "Mesma turma, ritmo compartilhado"),
            ("Grupo (3 a 5)", "R$ [900]/aluno", "Melhor custo por aluno"),
            ("Turma/escola (in company)", "Sob consulta", "Cachê fechado por turma")]
    t = doc.add_table(rows=1, cols=3)
    table_borders(t, 4, "D7DEEA"); set_width(t, 6.5, [0.34, 0.28, 0.38])
    h = t.rows[0].cells
    for i, txt in enumerate(["Modalidade", "Por aluno (10h)", "Observação"]):
        shade(h[i], BLUE); set_cell_margins(h[i], 55, 55, 110, 110)
        cell_text(h[i], txt, 9.5, WHITE, bold=True)
    for idx, (a, b, c_) in enumerate(rows):
        row = t.add_row().cells
        bg = WHITE if idx % 2 == 0 else LIGHT2
        for c in row:
            shade(c, bg); set_cell_margins(c, 55, 55, 110, 110); vcenter(c)
        cell_text(row[0], a, 10, NAVY, bold=True)
        cell_text(row[1], b, 10, GRAY)
        cell_text(row[2], c_, 9.5, GRAY)


def numbered(doc, items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph(); para_spacing(p, 1, 1, 1.16)
        p.paragraph_format.left_indent = Inches(0.18)
        rn = p.add_run(f"{i}.  "); runfmt(rn, 10.5, BLUE, bold=True)
        r = p.add_run(it); runfmt(r, 10.5, GRAY)


def build(logo_path=None, out="docs/cursos/Proposta_Comercial_MUN_AluizioEducacao.docx"):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(GRAY)

    brand_header(doc, "Proposta comercial · Mentoria 10 horas", logo_path)
    title_block(doc)
    info_block(doc)

    section_bar(doc, "1. Objeto")
    body(doc, "Programa de mentoria individual de 10 horas para preparação de simulações "
              "Model United Nations (MUN), com foco em compreender as estruturas da ONU e "
              "aplicá-las na resolução de problemas das simulações — negociação, redação de "
              "resoluções e mobilização/obtenção de recursos.")

    section_bar(doc, "2. O que o aluno alcança")
    for b in [
        "Identificar o órgão/comitê competente para cada tipo de problema.",
        "Dominar as regras de procedimento (moções, caucus, votação).",
        "Redigir position paper, working paper e draft resolution.",
        "Negociar, formar blocos e construir consenso.",
        "Mobilizar e obter fundos para suas propostas em comitê.",
        "Apresentar-se com postura e oratória diplomáticas.",
    ]:
        bullet(doc, b)

    spacer(doc, 3)
    section_bar(doc, "3. Como funciona")
    for label, txt in [
        ("Carga horária:", "10 horas, em 5 encontros de 2 horas."),
        ("Modalidade:", "[ ] Online ao vivo   [ ] Presencial — local: [ ]."),
        ("Frequência sugerida:", "1 encontro por semana (programa em ~5 semanas)."),
        ("Material incluído:", "modelos comentados de documentos, glossário ONU/MUN "
                               "(PT-BR/EN), guia de regras de procedimento e plano de estudos."),
        ("Acompanhamento:", "feedback por rubrica em cada encontro e plano de evolução final."),
    ]:
        bullet(doc, txt, label=label)
    spacer(doc, 1)
    callout(doc, "Conteúdo programático",
            "O detalhamento aula por aula está no documento “Planejamento — Geopolítica "
            "Aplicada às Simulações Model UN (10h)”.", BLUE)

    spacer(doc, 2)
    section_bar(doc, "4. Investimento", "valores de referência — ajustáveis")
    invest_table(doc)
    spacer(doc, 2)
    body(doc, "Formas de pagamento.", bold=True, after=1)
    bullet(doc, "À vista (PIX/transferência) com desconto, ou")
    bullet(doc, "Parcelado em até [3]x (cartão/boleto), sujeito a acréscimo da operadora.")
    spacer(doc, 3)
    body(doc, "Cenários alternativos (opcionais).", bold=True, after=2)
    cenarios_table(doc)

    spacer(doc, 3)
    section_bar(doc, "5. Condições")
    for b in [
        "Reserva confirmada após pagamento da 1ª parcela ou sinal de [50%].",
        "Remarcações com antecedência mínima de [24h]; faltas sem aviso não são repostas.",
        "Materiais de uso pessoal do aluno; não podem ser redistribuídos.",
        "Emissão de [recibo / nota fiscal] conforme combinado.",
    ]:
        bullet(doc, b)

    spacer(doc, 3)
    section_bar(doc, "6. Próximos passos")
    numbered(doc, [
        "Aprovação desta proposta por [e-mail/WhatsApp].",
        "Definição do calendário dos 5 encontros.",
        "Pagamento do sinal e envio dos materiais de boas-vindas.",
    ])

    spacer(doc, 8)
    callout(doc, "Aluizio Educação",
            "[e-mail]  ·  [telefone/WhatsApp]  ·  [site/redes]", NAVY)

    base.add_footer(doc)
    doc.save(out)
    print("Documento salvo em:", out)


if __name__ == "__main__":
    logo = sys.argv[1] if len(sys.argv) > 1 else None
    build(logo)
