#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera o CADERNO DE TRABALHO DO DELEGADO (Model UN) em Word (.docx) — artefato
prático e preenchível para o aluno trabalhar as temáticas do curso.

Uso:
    python3 scripts/gerar_kit_delegado_mun.py [caminho_logo.png]
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import gerar_planejamento_mun as base
from gerar_planejamento_mun import (
    NAVY, BLUE, GOLD, LIGHT, LIGHT2, GRAY, WHITE, FONT, DISPLAY,
    shade, set_cell_margins, vcenter, table_borders, no_borders, set_width,
    runfmt, para_spacing, cell_text, body, bullet, callout, spacer, section_bar,
)

LINE = "C9CFDB"


# ---------------------------------------------------------------- helpers extra
def header(doc, logo_path=None):
    t = doc.add_table(rows=1, cols=2)
    no_borders(t); set_width(t, 6.5, [0.62, 0.38])
    left, right = t.rows[0].cells
    shade(left, NAVY); shade(right, NAVY)
    set_cell_margins(left, 140, 140, 200, 120); set_cell_margins(right, 140, 140, 120, 200)
    vcenter(left); vcenter(right)
    left.text = ""
    p = left.paragraphs[0]; para_spacing(p, 0, 0, 1.0)
    ok = False
    if logo_path:
        try:
            p.add_run().add_picture(logo_path, height=Inches(0.55)); ok = True
        except Exception:
            ok = False
    if not ok:
        r = p.add_run("Aluizio"); runfmt(r, 22, WHITE, bold=True, font=DISPLAY)
        r2 = p.add_run("  Educação"); runfmt(r2, 22, GOLD, bold=True, font=DISPLAY)
        p2 = left.add_paragraph(); para_spacing(p2, 2, 0, 1.0)
        r3 = p2.add_run("Mentoria & preparação acadêmica"); runfmt(r3, 9, "C7D0E0", italic=True)
    right.text = ""
    pr = right.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT; para_spacing(pr, 0, 0, 1.0)
    r = pr.add_run("CADERNO DO DELEGADO"); runfmt(r, 10, GOLD, bold=True)
    pr2 = right.add_paragraph(); pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; para_spacing(pr2, 2, 0, 1.0)
    r = pr2.add_run("Model UN · material de trabalho"); runfmt(r, 9, "C7D0E0")


def title_block(doc):
    p = doc.add_paragraph(); para_spacing(p, 14, 0, 1.05)
    r = p.add_run("Caderno de Trabalho do Delegado"); runfmt(r, 22, NAVY, bold=True, font=DISPLAY)
    p2 = doc.add_paragraph(); para_spacing(p2, 2, 6, 1.15)
    r = p2.add_run("Ferramentas, modelos e exercícios preenchíveis para preparar e atuar "
                   "em simulações Model UN — da pesquisa do tema à obtenção de fundos.")
    runfmt(r, 11.5, GRAY, italic=True)
    line = doc.add_table(rows=1, cols=1); no_borders(line); set_width(line, 6.5, [1.0])
    shade(line.rows[0].cells[0], GOLD); cell_text(line.rows[0].cells[0], "", size=2)


def subhead(doc, text, before=8):
    p = doc.add_paragraph(); para_spacing(p, before, 2, 1.0)
    r = p.add_run(text); runfmt(r, 11, NAVY, bold=True)
    return p


def fill_label(doc, text):
    p = doc.add_paragraph(); para_spacing(p, 6, 0, 1.0)
    r = p.add_run(text); runfmt(r, 9.5, BLUE, bold=True)
    return p


def fill_lines(doc, n=2, before=8):
    """n linhas pontilhadas para escrever à mão/digitar."""
    for i in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before if i == 0 else 12)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'dotted'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), LINE)
        pBdr.append(bottom); pPr.append(pBdr)


def fill_box(doc, lines=3, label=None):
    """Caixa com borda leve para texto livre (altura ~ nº de linhas)."""
    if label:
        fill_label(doc, label)
    t = doc.add_table(rows=1, cols=1)
    table_borders(t, 4, LINE); set_width(t, 6.5, [1.0])
    c = t.rows[0].cells[0]
    shade(c, LIGHT2); set_cell_margins(c, 90, 90, 120, 120)
    c.text = ""
    for i in range(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        para_spacing(p, 0, 0, 1.6)
        p.add_run(" ")
    spacer(doc, 2)


def ref_table(doc, headers, rows, fracs, head_color=BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    table_borders(t, 4, "D7DEEA"); set_width(t, 6.5, fracs)
    h = t.rows[0].cells
    for i, txt in enumerate(headers):
        shade(h[i], head_color); set_cell_margins(h[i], 55, 55, 100, 100)
        cell_text(h[i], txt, 9.5, WHITE, bold=True)
    for idx, row in enumerate(rows):
        cells = t.add_row().cells
        bg = WHITE if idx % 2 == 0 else LIGHT2
        for j, c in enumerate(cells):
            shade(c, bg); set_cell_margins(c, 50, 50, 100, 100); vcenter(c)
            cell_text(c, row[j], 9.7, GRAY)
    spacer(doc, 2)


def numbered(doc, items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph(); para_spacing(p, 1, 1, 1.16)
        p.paragraph_format.left_indent = Inches(0.18)
        rn = p.add_run(f"{i}.  "); runfmt(rn, 10.5, BLUE, bold=True)
        r = p.add_run(it); runfmt(r, 10.5, GRAY)


def checklist(doc, items):
    for it in items:
        p = doc.add_paragraph(); para_spacing(p, 1, 1, 1.18)
        p.paragraph_format.left_indent = Inches(0.18)
        rb = p.add_run("☐  "); runfmt(rb, 11, BLUE, bold=True)
        r = p.add_run(it); runfmt(r, 10.5, GRAY)


# ------------------------------------------------------------------- conteúdo
def build(logo_path=None, out="docs/cursos/Caderno_Delegado_MUN_AluizioEducacao.docx"):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(GRAY)

    header(doc, logo_path)
    title_block(doc)

    # Como usar
    section_bar(doc, "Como usar este caderno")
    body(doc, "Este caderno acompanha o curso e serve para você praticar. Use as tabelas de "
              "referência durante o estudo e os modelos preenchíveis para montar seus "
              "documentos. As partes com linhas pontilhadas e caixas são para você completar.")

    # 1. Glossário
    section_bar(doc, "1. Glossário essencial ONU / MUN")
    ref_table(doc,
        ["Termo", "O que é"],
        [
         ["AGNU", "Assembleia Geral — deliberação universal; recomendações e orçamento."],
         ["CSNU", "Conselho de Segurança — paz e segurança; resoluções vinculantes; veto dos P5."],
         ["ECOSOC", "Conselho Econômico e Social — coordena temas econômicos, sociais e agências."],
         ["Caucus", "Debate informal — moderado (falas curtas por tema) ou não-moderado (negociação livre)."],
         ["GSL", "General Speakers' List — lista oficial de oradores do comitê."],
         ["Position paper", "Documento que apresenta a posição do país sobre o tema."],
         ["Draft resolution", "Projeto de resolução em negociação no comitê."],
         ["Cláusula operativa", "Item que determina ações concretas da resolução."],
         ["Trust fund", "Fundo fiduciário; recebe contribuições voluntárias com finalidade definida."],
        ],
        [0.22, 0.78])

    # 2. Mapa órgão x problema
    doc.add_page_break()
    section_bar(doc, "2. Qual órgão para qual problema?")
    body(doc, "Referência rápida para escolher o comitê competente — e espaço para você "
              "acrescentar os temas da sua simulação.")
    ref_table(doc,
        ["Tipo de problema", "Órgão / comitê indicado"],
        [
         ["Conflito armado, ameaça à paz", "Conselho de Segurança (CSNU)"],
         ["Tema global amplo, recomendação", "Assembleia Geral (AGNU)"],
         ["Crise humanitária / refugiados", "ECOSOC, ACNUR, OCHA, PMA"],
         ["Saúde / epidemia", "OMS (agência especializada)"],
         ["Desenvolvimento / pobreza", "PNUD, ECOSOC"],
         ["Infância", "UNICEF"],
         ["Disputa jurídica entre Estados", "Corte Internacional de Justiça (CIJ)"],
        ],
        [0.5, 0.5])
    fill_label(doc, "Acrescente os temas da sua simulação:")
    fill_lines(doc, 3)

    # 3. Cheat sheet de moções
    doc.add_page_break()
    section_bar(doc, "3. Cheat sheet — moções e pontos")
    ref_table(doc,
        ["Moção / ponto", "Quando usar"],
        [
         ["Abrir o debate", "No início, para começar a sessão."],
         ["Caucus moderado", "Debater um subtema com falas curtas e ordenadas."],
         ["Caucus não-moderado", "Negociar livremente e redigir documentos."],
         ["Encerrar o debate", "Quando o tema está maduro para votação."],
         ["Ponto de ordem", "Apontar erro de procedimento da mesa."],
         ["Ponto de informação", "Fazer pergunta (à mesa ou a um orador)."],
         ["Ponto de privilégio pessoal", "Resolver desconforto (não escuta, sala, etc.)."],
        ],
        [0.32, 0.68], head_color=GOLD)

    # 4. Pesquisa do país/tema
    doc.add_page_break()
    section_bar(doc, "4. Ficha de pesquisa — país e tema")
    fill_box(doc, 1, label="País / delegação representada:")
    fill_box(doc, 1, label="Comitê e tema da pauta:")
    fill_box(doc, 3, label="Posição oficial do país sobre o tema (3 pontos-chave):")
    fill_box(doc, 2, label="Aliados prováveis / blocos:")
    fill_box(doc, 2, label="Países com posição oposta:")
    fill_box(doc, 2, label="Dados e fontes (documentos da ONU, resoluções anteriores):")

    # 5. Position paper template
    doc.add_page_break()
    section_bar(doc, "5. Modelo de Position Paper")
    body(doc, "Estrutura sugerida — preencha cada parte.")
    fill_box(doc, 2, label="A) Contexto do tema (por que importa):")
    fill_box(doc, 3, label="B) Posição do país (o que defende e por quê):")
    fill_box(doc, 3, label="C) Propostas de solução (o que sugere ao comitê):")
    fill_box(doc, 2, label="D) Possíveis alianças e linhas de negociação:")

    # 6. Anatomia da resolução + verbos
    doc.add_page_break()
    section_bar(doc, "6. Anatomia da resolução e banco de verbos")
    body(doc, "A resolução tem cláusulas PREAMBULARES (contexto/justificativa, em itálico) e "
              "OPERATIVAS (ações, numeradas). Cada tipo usa verbos próprios.")
    ref_table(doc,
        ["Verbos preambulares", "Verbos operativos"],
        [
         ["Reafirmando, Reconhecendo", "Decide, Solicita"],
         ["Profundamente preocupado(a) com", "Recomenda, Encoraja"],
         ["Tendo examinado, Observando", "Convoca, Estabelece"],
         ["Guiado(a) pelos princípios de", "Apela para, Insta"],
         ["Levando em conta", "Autoriza, Solicita financiamento de"],
        ],
        [0.5, 0.5])

    # 7. Draft resolution template (com financiamento)
    doc.add_page_break()
    section_bar(doc, "7. Modelo de Draft Resolution")
    fill_box(doc, 1, label="Comitê / Tema / Patrocinadores (sponsors):")
    fill_label(doc, "Cláusulas preambulares (contexto — comece pelos verbos da seção 6):")
    fill_lines(doc, 3)
    fill_label(doc, "Cláusulas operativas (ações numeradas — inclua prazos e responsáveis):")
    fill_lines(doc, 4)
    callout(doc, "Cláusula de financiamento (modelo)",
            "“Solicita a criação de um fundo fiduciário voluntário, administrado por [agência], "
            "destinado a [objetivo], financiado por contribuições de Estados-membros e parceiros, "
            "com prestação de contas semestral ao [órgão].”", GOLD)

    # 8. Toolkit de financiamento
    doc.add_page_break()
    section_bar(doc, "8. Como obter fundos — passo a passo", "tema central do aluno")
    numbered(doc, [
        "Defina o objetivo e estime o custo da iniciativa.",
        "Identifique a fonte certa: orçamento regular, contribuições voluntárias, trust fund ou pooled fund (ex.: CERF).",
        "Escolha a agência/fundo gestor (PNUD, ACNUR, PMA, UNICEF, OMS...).",
        "Mapeie doadores plausíveis (Estados, blocos, fundações, setor privado).",
        "Redija a cláusula operativa que cria/solicita o financiamento.",
        "Preveja prestação de contas e monitoramento (quem fiscaliza e quando).",
    ])
    fill_box(doc, 2, label="Minha estratégia de financiamento para a simulação:")
    fill_box(doc, 2, label="Doadores/parceiros que vou abordar:")

    # 9. Mapa de blocos
    doc.add_page_break()
    section_bar(doc, "9. Mapa de blocos e aliados")
    ref_table(doc,
        ["Bloco / país", "Interesse principal", "Pontos de acordo comigo"],
        [["", "", ""], ["", "", ""], ["", "", ""], ["", "", ""]],
        [0.3, 0.4, 0.3])

    # 10. Discurso de abertura
    section_bar(doc, "10. Roteiro do discurso de abertura")
    body(doc, "Estrutura de 45–60 segundos: saudação à mesa → posição do país em 1 frase → "
              "1 a 2 propostas → chamado à cooperação.")
    fill_box(doc, 4, label="Meu discurso de abertura:")

    # 11. Checklist + 12. Rubrica
    doc.add_page_break()
    section_bar(doc, "11. Checklist pré-simulação")
    checklist(doc, [
        "Position paper pronto e revisado.",
        "Pesquisa do país e do tema concluída.",
        "Cheat sheet de moções à mão.",
        "Pelo menos 2 propostas de solução definidas.",
        "Estratégia de financiamento esboçada.",
        "Discurso de abertura ensaiado.",
        "Aliados e oposição mapeados.",
    ])

    spacer(doc, 4)
    section_bar(doc, "12. Rubrica de autoavaliação")
    body(doc, "Após cada simulação, dê nota de 1 a 5 e anote o que melhorar.")
    ref_table(doc,
        ["Critério", "Nota (1–5)", "O que melhorar"],
        [["Domínio das estruturas da ONU", "", ""],
         ["Uso correto do procedimento", "", ""],
         ["Qualidade dos documentos", "", ""],
         ["Negociação e construção de consenso", "", ""],
         ["Oratória e postura", "", ""],
         ["Estratégia de financiamento", "", ""]],
        [0.46, 0.16, 0.38])

    spacer(doc, 6)
    callout(doc, "Bom trabalho, delegado(a)!",
            "Releia este caderno antes de cada simulação. A diferença entre um bom delegado e "
            "um delegado excelente está na preparação — e em saber transformar a estrutura da "
            "ONU em ação concreta.", NAVY)

    base.add_footer(doc)
    doc.save(out)
    print("Documento salvo em:", out)


if __name__ == "__main__":
    logo = sys.argv[1] if len(sys.argv) > 1 else None
    build(logo)
