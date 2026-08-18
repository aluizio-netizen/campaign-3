# -*- coding: utf-8 -*-
"""
Gera o ROTEIRO do professor (Word) para a aula "Somos um em Cristo".
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1F, 0x3A, 0x5F)
DOURADO = RGBColor(0xB0, 0x8A, 0x10)
TEAL = RGBColor(0x2E, 0x6E, 0x6A)
CINZA = RGBColor(0x33, 0x33, 0x33)

doc = Document()

# margens
for sec in doc.sections:
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.font.color.rgb = CINZA


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def h_title(text, size=22, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER,
            after=4, bold=True, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = "Georgia"
    return p


def secao(num, titulo, tempo):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. {titulo}")
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = AZUL
    r.font.name = "Georgia"
    r2 = p.add_run(f"   ({tempo})")
    r2.font.size = Pt(12); r2.font.bold = True; r2.font.italic = True
    r2.font.color.rgb = DOURADO
    # linha inferior
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "C9A227")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def sub(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.bold = True; r.font.size = Pt(11.5)
    r.font.color.rgb = color
    return p


def slide_tag(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("  ▶ " + texto + "  ")
    r.font.size = Pt(9.5); r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # highlight
    rPr = r._r.get_or_add_rPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), "2E6E6A")
    rPr.append(sh)
    return p


def diga(texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("“" + texto + "”")
    r.font.italic = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def li(texto, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = AZUL
        r2 = p.add_run(texto); r2.font.size = Pt(11)
    else:
        r = p.add_run(texto); r.font.size = Pt(11)
    return p


def pergunta(texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("? ")
    r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = DOURADO
    r2 = p.add_run(texto)
    r2.font.size = Pt(11); r2.font.italic = True; r2.font.color.rgb = AZUL
    return p


# ===================== CABECALHO =====================
h_title("ROTEIRO DO PROFESSOR", 14, DOURADO, after=2)
h_title("SOMOS UM EM CRISTO", 28, AZUL, after=2)
h_title("Uma só oliveira, um só corpo, um só amor", 14, TEAL, after=10,
        bold=False, italic=True)

# ficha tecnica
t = doc.add_table(rows=5, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.style = "Table Grid"
ficha = [
    ("Tema", "Somos um em Cristo"),
    ("Texto-base", "Romanos 11 e 12 (em especial Rm 12.1-21)"),
    ("Versículo-chave", "“Conquanto muitos, somos um só corpo em Cristo, e membros uns dos outros” (Rm 12.5)"),
    ("Público", "Classe de Adultos — Escola Bíblica Dominical"),
    ("Igreja / Duração", "Segunda Igreja Batista de Campo Grande-MS  ·  60 minutos"),
]
for i, (k, v) in enumerate(ficha):
    c0 = t.rows[i].cells[0]; c1 = t.rows[i].cells[1]
    c0.width = Inches(1.8); c1.width = Inches(5.3)
    shade(c0, "1F3A5F")
    p0 = c0.paragraphs[0]; r0 = p0.add_run(k)
    r0.font.bold = True; r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r0.font.size = Pt(10.5)
    p1 = c1.paragraphs[0]; r1 = p1.add_run(v); r1.font.size = Pt(10.5)

doc.add_paragraph()

# OBJETIVO
sub("OBJETIVO DA AULA", AZUL)
p = doc.add_paragraph()
p.add_run("Ao final da aula, o aluno deverá compreender que, em Cristo, "
          "Deus forma um só povo de muitos povos, e ser desafiado a viver "
          "essa unidade na prática — guardando a unidade, valorizando a "
          "diversidade de dons, servindo uns aos outros e amando com "
          "sinceridade.").font.size = Pt(11)

# MATERIAIS / PREPARO
sub("PREPARO E MATERIAIS", AZUL)
li("Bíblia, projetor e a apresentação “Somos_um_em_Cristo_apresentacao.pptx” (22 slides).")
li("Leia antecipadamente Romanos 11 e 12 e ore pela classe.")
li("Tenha à mão giz/quadro ou folhas para a dinâmica de abertura.")

# DIVISAO DO TEMPO (tabela)
sub("VISÃO GERAL DO TEMPO (60 min)", AZUL)
tt = doc.add_table(rows=8, cols=3); tt.style = "Table Grid"
tt.alignment = WD_TABLE_ALIGNMENT.CENTER
tempo_rows = [
    ("Momento", "Conteúdo", "Tempo"),
    ("Abertura", "Acolhida, oração e quebra-gelo", "7 min"),
    ("Introdução", "Da doutrina à vida (Rm 1–11 → 12)", "6 min"),
    ("Ponto 1", "Um só povo: a oliveira de Deus (Rm 11)", "12 min"),
    ("Ponto 2", "Um só corpo em Cristo (Rm 12.1-8)", "15 min"),
    ("Ponto 3", "Um só amor (Rm 12.9-21)", "12 min"),
    ("Aplicação", "E na nossa igreja? + Conclusão", "6 min"),
    ("Encerramento", "Oração final e desafio da semana", "2 min"),
]
for i, row in enumerate(tempo_rows):
    cells = tt.rows[i].cells
    cells[0].width = Inches(1.3); cells[1].width = Inches(4.3); cells[2].width = Inches(0.9)
    for j, val in enumerate(row):
        pr = cells[j].paragraphs[0]
        run = pr.add_run(val)
        run.font.size = Pt(10)
        if i == 0:
            run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade(cells[j], "1F3A5F")
        else:
            if j == 0:
                run.font.bold = True; run.font.color.rgb = AZUL

doc.add_page_break()

# ===================== DESENVOLVIMENTO =====================
h_title("DESENVOLVIMENTO DA AULA", 16, AZUL, align=WD_ALIGN_PARAGRAPH.LEFT, after=8)

# --- ABERTURA ---
secao("1", "Abertura — acolhida e quebra-gelo", "7 min")
slide_tag("Slides 1 e 2 — Capa e versículo-tema")
sub("Acolha a classe e ore.")
sub("Quebra-gelo (dinâmica do corpo):")
diga("Quem consegue bater palmas usando uma só mão? E enxergar com o "
     "cotovelo? Cada parte do nosso corpo tem uma função — e nenhuma "
     "funciona sozinha. Hoje veremos que a igreja é assim: muitos "
     "membros, um só corpo.")
li("Leia juntos o versículo-tema (Rm 12.5) e anuncie o tema da aula.")
pergunta("O que vem à sua mente quando ouve a expressão “somos um em Cristo”?")

# --- INTRODUCAO ---
secao("2", "Introdução — da doutrina à vida", "6 min")
slide_tag("Slides 3 e 4 — Roteiro e introdução")
sub("Contexto de Romanos:")
li("Romanos é o maior tratado teológico do NT. Nos capítulos 1–11, "
   "Paulo expõe a doutrina; a partir do capítulo 12, trata da ética. "
   "Ele passa da teologia para a vida, do ensino para o dever.")
li("O capítulo 11 termina em adoração: “A Ele seja a glória para "
   "sempre! Amém” (Rm 11.36). A vida transformada desemboca na "
   "doxologia, e a unidade nasce da adoração.")
diga("Não podemos ter um relacionamento vertical correto com Deus se "
     "os relacionamentos horizontais com os irmãos estão errados.")
li("Apresente o roteiro dos três pontos: um só povo, um só corpo, um só amor.")

# --- PONTO 1 ---
secao("3", "Um só povo: a oliveira de Deus (Romanos 11)", "12 min")
slide_tag("Slides 5 a 8")

sub("a) Deus não rejeitou o seu povo (Rm 11.1-10)")
li("À pergunta “Deus rejeitou Israel?”, Paulo responde: “De modo "
   "nenhum!”. Ele mesmo é a prova — israelita, da descendência de Abraão.")
li("A rejeição é apenas parcial: sempre houve um remanescente fiel "
   "(como os 7.000 que não dobraram os joelhos a Baal). Esse "
   "remanescente existe “segundo a eleição da graça” (11.5) — não por "
   "obras, “do contrário, a graça já não seria graça”.")

sub("b) Gentios enxertados na oliveira (Rm 11.11-24)")
li("A queda de Israel abriu espaço para a salvação dos gentios, e isso "
   "deve provocar Israel a um santo ciúme, para a sua restauração.")
li("A figura da oliveira: a raiz é Abraão; os ramos naturais são os "
   "judeus; os gentios são ramos silvestres enxertados pela fé.")
li("Advertência contra o orgulho: “Não te glories contra os ramos” "
   "(11.18). O gentio não sustenta a raiz — a raiz é que o sustenta.")

sub("c) A verdade central: uma só oliveira (ênfase!)", DOURADO)
diga("Uma só oliveira representa todos os salvos, sem importar a sua "
     "origem. Para judeus e gentios, a salvação é a mesma: sobre a base "
     "da expiação de Cristo, pela graça, por meio da fé.")
li("Judeus e gentios foram chamados a formar “um só corpo em Cristo, a "
   "saber, a igreja” (Ef 3.4-6). Aqui está a ponte para Romanos 12.")
pergunta("Se a unidade do povo de Deus não é uniformidade cultural, "
         "o que nos torna um? (Resp.: a mesma raiz e a mesma graça em Cristo.)")

# --- PONTO 2 ---
secao("4", "Um só corpo em Cristo (Romanos 12.1-8)", "15 min")
slide_tag("Slides 9 a 15")

sub("A base: vidas transformadas (Rm 12.1-2)")
li("Antes de falar do corpo, Paulo trata do relacionamento com Deus: "
   "apresentar o corpo em sacrifício vivo, santo e agradável — nosso "
   "culto racional.")
li("“Não vos conformeis com este mundo, mas transformai-vos pela "
   "renovação da vossa mente.” A palavra grega evoca uma metamorfose: "
   "mudança de dentro para fora. Só pessoas transformadas vivem "
   "relacionamentos transformados.")

sub("As quatro verdades sobre o corpo (Rm 12.3-8)", DOURADO)
li("Unidade (12.5): ", "“Somos um só corpo.” Uma só família, um só "
   "rebanho. O que dá unidade é estarmos ligados à mesma Cabeça, Cristo, "
   "e irrigados pelo mesmo sangue. A unidade é recebida, não produzida.")
li("Diversidade (12.4-6a): ", "“Nem todos os membros têm a mesma "
   "função.” A marca das obras de Deus é a uniformidade dentro da "
   "unidade. A diversidade não ameaça — enriquece.")
li("Mutualidade (12.5): ", "“Somos membros uns dos outros.” Não "
   "competimos; servimos. Como as mãos levam o alimento à boca, cada "
   "membro cuida do todo (1Co 12.25).")
li("Utilidade (12.6-8): ", "os dons são dados pela Trindade para "
   "edificar o corpo, nunca para exibição: profecia, ministério, "
   "ensino, exortação, contribuição, presidência/liderança e "
   "misericórdia.")
pergunta("Você já descobriu o seu dom? Como o tem usado para edificar a igreja?")

# --- PONTO 3 ---
secao("5", "Um só amor (Romanos 12.9-21)", "12 min")
slide_tag("Slides 16 a 19")

sub("O amor que sustenta a unidade (Rm 12.9-12)")
li("O amor (ágape) é o sistema circulatório do corpo espiritual: faz "
   "todos os membros funcionarem de forma saudável.")
li("“O amor seja sem hipocrisia” (12.9) — amor sincero, sem máscara, "
   "com discernimento: detestar o mal, apegar-se ao bem.")
li("Afeição fraternal e honra mútua: “preferindo-vos em honra uns aos "
   "outros” (12.10).")

sub("Um amor que se abre (Rm 12.13-16)")
li("Coração aberto", " — amor cordial e fraterno;")
li("Lábios abertos", " — abençoar e não amaldiçoar (12.14);")
li("Mãos abertas", " — repartir com os santos nas necessidades (12.13a);")
li("Casa aberta", " — praticar a hospitalidade, acolher (12.13b).")

sub("Vencer o mal com o bem (Rm 12.17-21)")
li("A unidade cristã alcança até os inimigos: não retaliar, não criar "
   "conflitos desnecessários, não vingar-se, não guardar mágoa.")
li("“Se possível, quanto depender de vós, tende paz com todos os "
   "homens” (12.18). O cristão é pacificador.")
diga("Não te deixes vencer do mal, mas vence o mal com o bem (Rm 12.21).")
pergunta("Há algum relacionamento em que você precisa “vencer o mal "
         "com o bem” esta semana?")

# --- APLICACAO ---
secao("6", "Aplicação e conclusão", "6 min")
slide_tag("Slides 20 e 21")
sub("Perguntas de aplicação (escolha 2 ou 3 conforme o tempo):")
pergunta("Guardo a unidade ou alimento divisões e fofocas? (12.16)")
pergunta("Conheço e uso o meu dom — ou vivo como espectador?")
pergunta("Valorizo a diversidade, ou exijo que todos sejam iguais a mim?")
pergunta("Minha casa, mãos e lábios estão abertos aos irmãos?")
sub("Síntese final:")
diga("Em Cristo, Deus fez de muitos um só povo: uma só oliveira, um só "
     "corpo, um só amor. A unidade é dom recebido pela graça; a "
     "diversidade nos enriquece; e o amor mantém tudo unido. Ser um em "
     "Cristo não é um ideal distante — é uma realidade para viver todos "
     "os dias.")

# --- ENCERRAMENTO ---
secao("7", "Encerramento", "2 min")
slide_tag("Slide 22")
li("Conduza a oração final, agradecendo pela unidade do corpo de Cristo "
   "e pedindo graça para vivê-la.")
sub("Desafio da semana:")
diga("Procure, durante esta semana, servir um irmão com o seu dom e "
     "buscar a paz com alguém de quem você esteja afastado.")

# RODAPE final
doc.add_paragraph()
pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run("“A Ele seja a glória para sempre! Amém.” (Romanos 11.36)")
rf.font.italic = True; rf.font.size = Pt(11); rf.font.color.rgb = AZUL

doc.save("/home/user/campaign-3/aula-somos-um-em-cristo/Somos_um_em_Cristo_roteiro.docx")
print("Roteiro DOCX gerado.")
